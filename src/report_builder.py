"""Generate the Word department report.

Ported from build_report.py, parameterized by year and operating on an
in-memory results dict + a directory of pre-generated chart PNGs.
"""
from __future__ import annotations

import os
from io import BytesIO
from typing import Union

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

NAVY = RGBColor(0x1F, 0x4E, 0x79)
TEAL = RGBColor(0x2E, 0x8B, 0x8B)
GRAY = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0xC0, 0x50, 0x4D)


def build_report(
    results: dict,
    charts_dir: str,
    year: Union[int, str],
    output_path: Union[str, BytesIO],
) -> None:
    """Generate the Word report and write to a path or BytesIO buffer."""
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    RES = results
    CHARTS = charts_dir

    # ----- helpers -----
    def add_heading(text, level=1, color=NAVY):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = color
        if level == 1:
            run.font.size = Pt(18)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
        elif level == 2:
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        else:
            run.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
        return p

    def add_para(text, bold=False, italic=False, size=11, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        p.paragraph_format.space_after = Pt(4)
        return p

    def shade_cell(cell, fill):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        tcPr.append(shd)

    def make_table(headers, rows, col_widths=None, header_fill="1F4E79"):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ""
            p = hdr[i].paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            shade_cell(hdr[i], header_fill)
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = t.rows[r_i + 1].cells[c_i]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(val) if val is not None else "—")
                run.font.size = Pt(10)
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in t.rows:
                    row.cells[i].width = w
        return t

    def add_image(name, width_in=6.5):
        path = os.path.join(CHARTS, f"{name}.png")
        if not os.path.exists(path):
            add_para(f"[Chart not available: {name}]", italic=True, size=9, color=GRAY)
            return
        doc.add_picture(path, width=Inches(width_in))
        p = doc.paragraphs[-1]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)

    def fmt(v, spec=""):
        if v is None:
            return "—"
        try:
            if spec:
                return format(v, spec)
            return str(v)
        except Exception:
            return str(v)

    # ----- TITLE PAGE -----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("VUMC Department of Anesthesiology")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"Faculty Engagement & Well-being Survey — {year}")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = TEAL

    date = doc.add_paragraph()
    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = date.add_run("Comprehensive Results Report")
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = GRAY

    doc.add_paragraph()

    # ----- EXECUTIVE SUMMARY -----
    add_heading("Executive Summary", 1)
    n_total = RES["sample"]["n_total"]
    n_clin = RES["sample"]["clinical"]
    n_res = RES["sample"]["research"]

    mz_c = RES["miniz"]["clinical"]
    wb = RES["wellbeing"]["overall"]
    nps = RES["nps"]["overall"]
    ret = RES["retention"]["overall"]

    add_para(
        f"A total of {n_total} faculty responded ({n_clin} clinical, {n_res} research, with the "
        f"remainder having missing clinical-care designation). Headline indicators are shown below; "
        f"detail and division-level breakdowns follow."
    )

    has_clin_total = "total" in mz_c
    tile_rows = []
    if has_clin_total:
        tile_rows.append([
            "MINI-Z total (clinical, ≥40 = joyful)",
            f"{mz_c['total']['mean']:.1f} / 50",
            f"{mz_c['total']['joyful_pct']:.1f}% joyful",
        ])
    tile_rows += [
        ["Burnout (clinical, definitely burning out or worse)",
         f"{fmt(mz_c['burnout_pct'], '.1f')}%",
         f"{mz_c['burnout_n']}/{mz_c['burnout_total']}"],
        ["High stress at work (clinical)",
         f"{fmt(mz_c['stress_pct'], '.1f')}%",
         f"{mz_c['stress_n']}/{mz_c['burnout_total']}"],
        ["Job satisfaction — agree/strongly agree (clinical)",
         f"{fmt(mz_c['jobsat_pct'], '.1f')}%",
         f"{mz_c['jobsat_n']}/{mz_c['burnout_total']}"],
    ]
    if wb:
        tile_rows.append(["Well-being (overall, 0–10)", f"{wb['mean']:.2f}",
                          f"High WBI (8–10): {wb['high_pct']:.1f}%"])
    if nps:
        tile_rows.append(["Net Promoter Score (overall)", f"{nps['nps']:.1f}",
                          f"Promoters {nps['promoters_pct']:.0f}% / Detractors {nps['detractors_pct']:.0f}%"])
    if ret:
        tile_rows.append(["Likely / very likely to still be at VUMC in 3 yr (overall)",
                          f"{ret['likely_stay_pct']:.1f}%",
                          f"At-risk (≤ neutral): {ret['at_risk_pct']:.1f}%"])

    make_table(["Indicator", "Result", "Detail"], tile_rows,
               col_widths=[Inches(3.4), Inches(1.6), Inches(2.0)])

    doc.add_paragraph()
    add_para("Key takeaways (auto-generated headline numbers):", bold=True)

    bullets = []
    if has_clin_total:
        bullets.append(
            f"Clinical MINI-Z total averages {mz_c['total']['mean']:.1f}/50, "
            f"with {mz_c['total']['joyful_pct']:.1f}% meeting the joyful threshold (≥40)."
        )
    bullets.append(
        f"Clinical faculty: {fmt(mz_c['burnout_pct'], '.1f')}% burnout, "
        f"{fmt(mz_c['stress_pct'], '.1f')}% high stress, "
        f"{fmt(mz_c['jobsat_pct'], '.1f')}% job satisfaction."
    )
    if wb:
        bullets.append(
            f"Well-being index (0–10) averages {wb['mean']:.2f} overall; "
            f"{wb['high_pct']:.1f}% high (8–10), {wb['low_pct']:.1f}% low (0–5)."
        )
    if nps:
        nps_c = RES['nps'].get('clinical') or {}
        nps_r = RES['nps'].get('research') or {}
        if nps_c and nps_r:
            bullets.append(
                f"NPS = {nps['nps']:.1f} overall (clinical {nps_c.get('nps','—')}, "
                f"research {nps_r.get('nps','—')})."
            )
        else:
            bullets.append(f"NPS = {nps['nps']:.1f} overall.")
    if ret:
        bullets.append(
            f"Retention risk (≤ neutral on intent to still be at VUMC in 3 yrs) "
            f"is {ret['at_risk_pct']:.1f}% overall."
        )
    if RES["retention"]["stay_reasons"]["items"]:
        top_drivers = [it["reason"] for it in RES["retention"]["stay_reasons"]["items"][:5]]
        bullets.append("Top stay-decision drivers: " + "; ".join(top_drivers) + ".")
    if RES["factors"]:
        top_factors = [f["factor"] for f in RES["factors"] if f["mean"] is not None][:3]
        bot_factors = [f["factor"] for f in sorted(RES["factors"], key=lambda x: x["mean"] or 99)
                       if f["mean"] is not None][:3]
        if top_factors:
            bullets.append("Strongest job-factor satisfiers: " + "; ".join(top_factors) + ".")
        if bot_factors:
            bullets.append("Largest job-factor dissatisfiers: " + "; ".join(bot_factors) + ".")

    for line in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(line).font.size = Pt(11)

    add_image("key_outcomes_by_group", 6.5)
    add_para("Figure 1. Headline outcomes — Clinical vs. Research faculty.",
             italic=True, size=10, color=GRAY)

    doc.add_page_break()

    # ----- METHODS NOTE -----
    add_heading("Methods Note", 1)
    add_para(
        "Data source: REDCap survey export (raw coded values). The test record (any response "
        "containing 'this is a test' in the free-text fields) was excluded. "
        f"All remaining respondents (n = {n_total}) are retained. 'Clinical' = answered Yes to "
        f"'Do you provide clinical care?' (n = {n_clin}); 'Research' = answered No (n = {n_res}). "
        "MINI-Z 2.0 scoring: each item scored 1–5 with 5 = best (per the official scoring manual); "
        "overall total is the sum of all 10 items (≥40 = joyful). "
        "Subscale 1 (Supportive Work Environment / Satisfaction) = Q1, Q2, Q3, Q4, Q9; "
        "Subscale 2 (Work Pace / EMR Stress) = Q5, Q6, Q7, Q8, Q10. Each subscale ranges 5–25; "
        "≥20 = highly supportive / reasonable pace, per the Mini-Z 2.0 manual (Linzer / InstitutePHI). "
        "'Burnout' = response 'definitely burning out' or worse on Q2 (raw ≤ 3); "
        "'High stress' = agree/strongly agree to 'A great deal of stress' on Q5 (raw ≤ 2); "
        "'Job satisfaction' = agree/strongly agree on Q1 (raw ≥ 4). Research faculty did not see "
        "clinical-specific items (Q4, Q6, Q7, Q10), so their MINI-Z subscales/total are not reported. "
        "Well-being index uses the full 0–10 Cantril-style item; bands are 0–5 low, 6–7 mid, 8–10 high. "
        "NPS = %Promoters (9–10) − %Detractors (0–6). Retention risk uses the 5-pt "
        "'In 3 years, I will still be working at VUMC' item; 'at risk' = neutral or worse. "
        "Job-factor scores rescaled from −2 (very negative impact) to +2 (very positive); 'N/A' "
        "coded missing. Leadership items normalized to 1–5 (5 = strongly agree)."
    )

    doc.add_page_break()

    # ----- DEMOGRAPHICS -----
    add_heading("1. Sample & Demographics", 1)
    add_para(f"Total respondents (test record excluded): n = {n_total}. "
             f"Clinical faculty: n = {n_clin}. Research faculty: n = {n_res}.")

    dem = RES["demographics"]["tables"]

    def dem_table(label, key):
        if key not in dem:
            return
        add_heading(label, 3)
        rows = [[d["value"], d["n"], f"{d['pct']:.1f}%"] for d in dem[key]]
        make_table(["Category", "n", "%"], rows,
                   col_widths=[Inches(4.0), Inches(1.0), Inches(1.0)])
        doc.add_paragraph()

    dem_table("Gender", "Gender")
    dem_table("Ethnicity", "Ethnicity")

    if RES["demographics"].get("race"):
        add_heading("Race (multi-select)", 3)
        rows = [[d["value"], d["n"], f"{d['pct']:.1f}%"] for d in RES["demographics"]["race"]]
        make_table(["Category", "n", "%"], rows,
                   col_widths=[Inches(4.0), Inches(1.0), Inches(1.0)])
        doc.add_paragraph()

    dem_table("Time at VUMC", "Time at VUMC")
    dem_table("FTE", "FTE")
    dem_table("Academic days", "Academic days FY")

    if "Primary Division" in dem:
        add_heading("Primary Division", 3)
        rows = [[d["value"], d["n"], f"{d['pct']:.1f}%"] for d in dem["Primary Division"]]
        make_table(["Division", "n", "%"], rows,
                   col_widths=[Inches(4.0), Inches(1.0), Inches(1.0)])
        doc.add_paragraph()
        add_image("division_count", 6.0)
        add_para("Figure 2. Respondents by primary division.",
                 italic=True, size=10, color=GRAY)

    doc.add_page_break()

    # ----- MINI-Z -----
    add_heading("2. MINI-Z (Burnout & Practice Joy)", 1)
    add_para(
        "MINI-Z 2.0 — 10 items, 1 = least favorable, 5 = most favorable. Item-level results below; "
        "subscales and overall scores follow. Research faculty did not receive items 4, 6, 7, 10, "
        "so their subscale/total scores are not reported."
    )

    def miniz_item_table(group_label, items):
        add_heading(group_label, 3)
        rows = []
        for it in items:
            if it["n"] == 0 or it["mean"] is None:
                rows.append([it["item"], "—", "—", "—", "—"])
            else:
                rows.append([
                    it["item"], it["n"], f"{it['mean']:.2f}",
                    f"{it['sd']:.2f}", f"{it['pct_favorable']:.1f}%",
                ])
        make_table(["Item", "n", "Mean", "SD", "% favorable (4–5)"], rows,
                   col_widths=[Inches(3.7), Inches(0.6), Inches(0.8), Inches(0.7), Inches(1.2)])
        doc.add_paragraph()

    miniz_item_table("Item-level — Overall", RES["miniz"]["overall"]["items"])
    miniz_item_table("Item-level — Clinical faculty", RES["miniz"]["clinical"]["items"])
    miniz_item_table("Item-level — Research faculty", RES["miniz"]["research"]["items"])

    add_image("miniz_items", 6.5)
    add_para("Figure 3. MINI-Z item means (clinical) — higher = better on all items.",
             italic=True, size=10, color=GRAY)

    if has_clin_total:
        add_heading("Subscales and Total — Clinical faculty (MINI-Z 2.0)", 2)
        mz = RES["miniz"]["clinical"]
        sub_rows = [
            ["Subscale 1 — Supportive Work Environment (Q1, Q2, Q3, Q4, Q9; range 5–25)",
             mz["subscale1"]["n"], f"{mz['subscale1']['mean']:.2f}",
             f"{mz['subscale1']['sd']:.2f}",
             f"{mz['subscale1']['supportive_pct']:.1f}% ≥20"],
            ["Subscale 2 — Work Pace / EMR Stress (Q5, Q6, Q7, Q8, Q10; range 5–25)",
             mz["subscale2"]["n"], f"{mz['subscale2']['mean']:.2f}",
             f"{mz['subscale2']['sd']:.2f}",
             f"{mz['subscale2']['lowstress_pct']:.1f}% ≥20"],
            ["MINI-Z Total (all 10 items; range 10–50)",
             mz["total"]["n"], f"{mz['total']['mean']:.2f}",
             f"{mz['total']['sd']:.2f}",
             f"{mz['total']['joyful_pct']:.1f}% ≥40 (joyful)"],
        ]
        make_table(["Score", "n", "Mean", "SD", "Manual threshold"], sub_rows,
                   col_widths=[Inches(3.5), Inches(0.5), Inches(0.8), Inches(0.7), Inches(1.5)])
        doc.add_paragraph()
        add_para(
            f"Clinical burnout: {mz['burnout_n']}/{mz['burnout_total']} = "
            f"{fmt(mz['burnout_pct'], '.1f')}%. "
            f"High stress: {mz['stress_n']}/{mz['burnout_total']} = "
            f"{fmt(mz['stress_pct'], '.1f')}%. "
            f"Job satisfaction: {mz['jobsat_n']}/{mz['burnout_total']} = "
            f"{fmt(mz['jobsat_pct'], '.1f')}%."
        )

    add_heading("Research faculty — MINI-Z items reported (no total/subscales due to skip logic)", 3)
    mr = RES["miniz"]["research"]
    add_para(
        f"Burnout: {mr['burnout_n']}/{mr['burnout_total']} = "
        f"{fmt(mr['burnout_pct'], '.1f')}%; "
        f"High stress: {mr['stress_n']}/{mr['burnout_total']} = "
        f"{fmt(mr['stress_pct'], '.1f')}%; "
        f"Job satisfaction: {mr['jobsat_n']}/{mr['burnout_total']} = "
        f"{fmt(mr['jobsat_pct'], '.1f')}%."
    )

    add_heading("MINI-Z by division (clinical-attached divisions; research listed for reference)", 2)
    rows = []
    for d in sorted(RES["miniz"]["by_division"],
                    key=lambda x: -((x.get("total") or {}).get("mean") or 0)):
        rows.append([
            d["division"], d["n"],
            f"{d['total']['mean']:.1f}" if d.get("total") and d["total"]["mean"] is not None else "—",
            f"{d['total']['joyful_pct']:.1f}%" if d.get("total") and d["total"]["joyful_pct"] is not None else "—",
            f"{d['subscale1']['mean']:.1f}" if d.get("subscale1") and d["subscale1"]["mean"] is not None else "—",
            f"{d['subscale2']['mean']:.1f}" if d.get("subscale2") and d["subscale2"]["mean"] is not None else "—",
            f"{fmt(d.get('burnout_pct'), '.1f')}%",
            f"{fmt(d.get('stress_pct'), '.1f')}%",
        ])
    make_table(["Division", "n", "Total mean", "% joyful (≥40)", "Subscale 1 (Sat)",
                "Subscale 2 (Stress)", "% burnout", "% high stress"], rows,
               col_widths=[Inches(1.4), Inches(0.5), Inches(0.8), Inches(0.9),
                           Inches(0.9), Inches(0.9), Inches(0.7), Inches(0.9)])
    doc.add_paragraph()
    add_image("miniz_by_division", 6.5)
    add_para("Figure 4. MINI-Z total and burnout/stress prevalence by division.",
             italic=True, size=10, color=GRAY)

    doc.add_page_break()

    # ----- WELL-BEING -----
    add_heading("3. Well-being Index (0–10)", 1)
    wb_o = RES["wellbeing"]["overall"]
    wb_c = RES["wellbeing"]["clinical"]
    wb_r = RES["wellbeing"]["research"]
    rows = []
    for label, w in [("Overall", wb_o), ("Clinical", wb_c), ("Research", wb_r)]:
        if w is None:
            rows.append([label, 0, "—", "—", "—", "—", "—"])
        else:
            rows.append([label, w["n"], f"{w['mean']:.2f}", f"{w['sd']:.2f}",
                         f"{w['low_pct']:.1f}%", f"{w['mid_pct']:.1f}%", f"{w['high_pct']:.1f}%"])
    make_table(["Group", "n", "Mean", "SD", "% Low (0–5)", "% Mid (6–7)", "% High (8–10)"], rows,
               col_widths=[Inches(1.0), Inches(0.6), Inches(0.8), Inches(0.7),
                           Inches(1.0), Inches(1.0), Inches(1.0)])
    doc.add_paragraph()

    add_heading("Well-being by division", 2)
    rows = []
    for d in sorted(RES["wellbeing"]["by_division"], key=lambda x: -x["mean"]):
        rows.append([d["division"], d["n"], f"{d['mean']:.2f}",
                     f"{d['low_pct']:.1f}%", f"{d['mid_pct']:.1f}%", f"{d['high_pct']:.1f}%"])
    make_table(["Division", "n", "Mean", "% Low", "% Mid", "% High"], rows,
               col_widths=[Inches(1.6), Inches(0.6), Inches(0.8),
                           Inches(1.0), Inches(1.0), Inches(1.0)])
    doc.add_paragraph()
    add_image("wbi_by_division", 6.5)
    add_para("Figure 5. Well-being index by division (mean and % high).",
             italic=True, size=10, color=GRAY)

    doc.add_page_break()

    # ----- NPS -----
    add_heading("4. Net Promoter Score (NPS)", 1)
    nps_o = RES["nps"]["overall"]
    nps_c = RES["nps"]["clinical"]
    nps_r = RES["nps"]["research"]
    rows = []
    for label, n in [("Overall", nps_o), ("Clinical", nps_c), ("Research", nps_r)]:
        if n is None:
            rows.append([label, 0, "—", "—", "—", "—", "—"])
            continue
        rows.append([label, n["n"], f"{n['mean']:.2f}",
                     f"{n['promoters_pct']:.1f}%", f"{n['passives_pct']:.1f}%",
                     f"{n['detractors_pct']:.1f}%", f"{n['nps']:.1f}"])
    make_table(["Group", "n", "Mean (0–10)", "Promoters (9–10)",
                "Passives (7–8)", "Detractors (0–6)", "NPS"], rows,
               col_widths=[Inches(0.9), Inches(0.5), Inches(0.9), Inches(1.3),
                           Inches(1.1), Inches(1.3), Inches(0.7)])
    doc.add_paragraph()

    add_heading("NPS by division", 2)
    rows = []
    for d in RES["nps"]["by_division"]:
        rows.append([d["division"], d["n"], f"{d['mean']:.2f}",
                     f"{d['promoters_pct']:.0f}%", f"{d['passives_pct']:.0f}%",
                     f"{d['detractors_pct']:.0f}%", f"{d['nps']:.1f}"])
    make_table(["Division", "n", "Mean", "Promoters", "Passives", "Detractors", "NPS"], rows,
               col_widths=[Inches(1.5), Inches(0.5), Inches(0.7), Inches(1.0),
                           Inches(1.0), Inches(1.0), Inches(0.7)])
    doc.add_paragraph()
    add_image("nps_by_division", 6.5)
    add_para("Figure 6. NPS by division.", italic=True, size=10, color=GRAY)

    doc.add_page_break()

    # ----- RETENTION -----
    add_heading("5. Retention & Risk of Leaving Vanderbilt", 1)
    ro = RES["retention"]["overall"]
    rc = RES["retention"]["clinical"]
    rr = RES["retention"]["research"]
    add_para("Item: 'In 3 years, I will still be working at VUMC.' (1 = very unlikely … 5 = very likely)")
    rows = []
    for label, x in [("Overall", ro), ("Clinical", rc), ("Research", rr)]:
        if x is None:
            rows.append([label, 0, "—", 0, 0, 0, 0, 0, "—", "—"])
            continue
        rows.append([label, x["n"], f"{x['mean']:.2f}",
                     x["vunlikely"], x["unlikely"], x["neutral"], x["likely"], x["vlikely"],
                     f"{x['at_risk_pct']:.1f}%", f"{x['likely_stay_pct']:.1f}%"])
    make_table(["Group", "n", "Mean", "VUnlik", "Unlik", "Neutral", "Likely", "VLikely",
                "% at-risk", "% likely stay"], rows,
               col_widths=[Inches(0.8), Inches(0.4), Inches(0.5), Inches(0.55),
                           Inches(0.5), Inches(0.65), Inches(0.5), Inches(0.55),
                           Inches(0.85), Inches(1.0)])
    doc.add_paragraph()

    add_heading("Retention risk by division (sorted high → low % likely-to-stay)", 2)
    rows = []
    for d in sorted(RES["retention"]["by_division"], key=lambda x: -x["likely_stay_pct"]):
        rows.append([d["division"], d["n"], f"{d['mean']:.2f}",
                     f"{d['at_risk_pct']:.1f}%", f"{d['likely_stay_pct']:.1f}%"])
    make_table(["Division", "n", "Mean", "% at-risk", "% likely stay"], rows,
               col_widths=[Inches(1.7), Inches(0.6), Inches(0.8), Inches(1.1), Inches(1.3)])
    doc.add_paragraph()
    add_image("retention_by_division", 6.5)
    add_para("Figure 7. Retention by division — % likely-to-stay (3-year horizon).",
             italic=True, size=10, color=GRAY)

    add_heading("Top drivers of stay-decision (multi-select)", 2)
    sr = RES["retention"]["stay_reasons"]
    add_para(f"Denominator: n = {sr['denom']} respondents to the retention block.")
    rows = [[it["reason"], it["n"], f"{it['pct']:.1f}%"] for it in sr["items"]]
    make_table(["Stay-decision driver", "n", "% of respondents"], rows,
               col_widths=[Inches(4.5), Inches(0.7), Inches(1.4)])
    doc.add_paragraph()
    add_image("stay_drivers", 6.5)
    add_para("Figure 8. Top stay-decision drivers (top 8 by % selecting).",
             italic=True, size=10, color=GRAY)

    le = RES["retention"].get("lifeevent")
    fte = RES["retention"].get("fte_reduce")
    if fte and fte["distribution"]:
        add_heading("Have you reduced or are you considering reducing your FTE?", 3)
        rows = [[d["value"], d["n"], f"{d['pct']:.1f}%"] for d in fte["distribution"]]
        make_table(["Response", "n", "%"], rows,
                   col_widths=[Inches(2.5), Inches(0.7), Inches(1.0)])
        doc.add_paragraph()
    if le and le.get("n"):
        add_para(
            f"Considering a major life event that may affect retention: "
            f"Yes = {le['yes_n']} ({le['yes_pct']:.1f}%); "
            f"Unsure = {le['unsure_n']} ({le['unsure_pct']:.1f}%); "
            f"No = {le['no_n']} ({le['no_pct']:.1f}%) of n = {le['n']}."
        )

    doc.add_page_break()

    # ----- JOB FACTORS -----
    add_heading("6. Job Factors — Most & Least Satisfaction", 1)
    add_para(
        "Each factor was rated for impact on job satisfaction on a 5-pt scale: very negative "
        "impact (−2), somewhat negative (−1), neutral (0), somewhat positive (+1), very positive "
        "(+2). 'Net positive' = % positive minus % negative. N varies (item N/A coded missing)."
    )
    factors = sorted([f for f in RES["factors"] if f["mean"] is not None],
                     key=lambda x: -x["mean"])

    add_heading("All factors ranked by mean impact", 2)
    rows = []
    for f in factors:
        rows.append([f["factor"], f["n"], f"{f['mean']:.2f}",
                     f"{f['pos_pct']:.1f}%", f"{f['neg_pct']:.1f}%",
                     f"{f['net_positive']:.1f}%"])
    make_table(["Job factor", "n", "Mean (−2…+2)", "% positive", "% negative", "Net positive"],
               rows,
               col_widths=[Inches(3.6), Inches(0.5), Inches(1.0),
                           Inches(0.9), Inches(0.9), Inches(0.8)])
    doc.add_paragraph()
    add_image("factors_top_bottom", 6.5)
    add_para("Figure 9. Top and bottom job-factor drivers of satisfaction.",
             italic=True, size=10, color=GRAY)

    doc.add_page_break()

    # ----- LEADERSHIP -----
    add_heading("7. Leadership Ratings", 1)
    add_para("All items use a 5-pt agreement scale (5 = strongly agree). "
             "'Agree' = % responding agree or strongly agree.")

    def lead_table(label, items):
        add_heading(label, 3)
        rows = []
        for it in items:
            if it["n"] == 0 or it["mean"] is None:
                rows.append([it["item"], 0, "—", "—", "—"])
            else:
                rows.append([it["item"], it["n"], f"{it['mean']:.2f}",
                             f"{it['sd']:.2f}", f"{it['agree_pct']:.1f}%"])
        make_table(["Item", "n", "Mean", "SD", "% agree"], rows,
                   col_widths=[Inches(4.3), Inches(0.5), Inches(0.7),
                               Inches(0.6), Inches(0.9)])
        doc.add_paragraph()

    lead_table("Overall (all faculty)", RES["leadership"]["overall"])
    lead_table("Clinical faculty", RES["leadership"]["clinical"])
    lead_table("Research faculty", RES["leadership"]["research"])

    add_image("leadership_overall", 6.5)
    add_para("Figure 10. Overall leadership ratings — mean by leadership tier.",
             italic=True, size=10, color=GRAY)

    add_heading("Leadership ratings by division (% agree on key items)", 2)
    add_para(
        "Items: DivCh Trust = trust my division chief; DivCh Lead = leads effectively; "
        "Chair Trust = trust dept chair; Chair Lead = leads effectively; "
        "VC Trust = trust VCs; VC Lead = exec leadership leads effectively; "
        "Community = sense of community/camaraderie; Account. = accountability.",
        size=10,
    )
    rows = []
    for d in RES["leadership"]["by_division"]:
        rows.append([
            d["division"], d["n"],
            fmt(d.get("trust1_agree_pct"), ".0f") + "%" if d.get("trust1_agree_pct") is not None else "—",
            fmt(d.get("lead1_agree_pct"), ".0f") + "%" if d.get("lead1_agree_pct") is not None else "—",
            fmt(d.get("trust3a_agree_pct"), ".0f") + "%" if d.get("trust3a_agree_pct") is not None else "—",
            fmt(d.get("lead3_agree_pct"), ".0f") + "%" if d.get("lead3_agree_pct") is not None else "—",
            fmt(d.get("trust2_agree_pct"), ".0f") + "%" if d.get("trust2_agree_pct") is not None else "—",
            fmt(d.get("lead2_agree_pct"), ".0f") + "%" if d.get("lead2_agree_pct") is not None else "—",
            fmt(d.get("community_agree_pct"), ".0f") + "%" if d.get("community_agree_pct") is not None else "—",
            fmt(d.get("accountability_agree_pct"), ".0f") + "%" if d.get("accountability_agree_pct") is not None else "—",
        ])
    make_table(
        ["Division", "n", "DivCh Trust", "DivCh Lead", "Chair Trust", "Chair Lead",
         "VC Trust", "VC Lead", "Community", "Account."],
        rows,
        col_widths=[Inches(1.0), Inches(0.4), Inches(0.7), Inches(0.7),
                    Inches(0.7), Inches(0.7), Inches(0.6), Inches(0.6),
                    Inches(0.7), Inches(0.7)],
    )
    doc.add_paragraph()

    # ----- CLOSING -----
    add_para(f"Report generated from REDCap raw-data export ({year} survey).",
             italic=True, size=9, color=GRAY)

    doc.save(output_path)
