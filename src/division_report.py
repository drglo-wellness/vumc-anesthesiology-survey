"""Per-division Word reports.

Generates a focused 3-5 page report for a single division and provides a helper
that loops over all divisions and bundles them into a ZIP for download.
"""
from __future__ import annotations

import io
import os
import re
import zipfile
from typing import Iterable, Union

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from src.analysis import (
    DIVISION_MAP,
    FACTOR_LABELS,
    LEADERSHIP_ITEMS,
)

# Map of leadership column name -> short display label, derived from LEADERSHIP_ITEMS
LEADERSHIP_LABELS = {col: lab for col, lab, _scale in LEADERSHIP_ITEMS}

NAVY = RGBColor(0x1F, 0x4E, 0x79)
TEAL = RGBColor(0x2E, 0x8B, 0x8B)
GRAY = RGBColor(0x55, 0x55, 0x55)
ACCENT = RGBColor(0xC0, 0x50, 0x4D)
LIGHT = "F4F6FA"
GOOD_FILL = "E2F0E5"   # light green
WARN_FILL = "FCE9CF"   # light orange
BAD_FILL = "F8D8D6"    # light red


# ---------------------------------------------------------------------------
# Helpers reused across multiple division reports
# ---------------------------------------------------------------------------

def _fmt_delta(x):
    if x is None:
        return "—"
    sign = "+" if x >= 0 else ""
    return f"{sign}{x:.2f}"


def _fmt_pct_delta(x):
    if x is None:
        return "—"
    sign = "+" if x >= 0 else ""
    return f"{sign}{x:.1f}pp"


def _compute_division_factors(df: pd.DataFrame) -> list[dict]:
    """Compute factor mean impacts for a single-division DataFrame.

    Filters the 8888 "N/A" REDCap code to NaN before computing means.
    """
    out = []
    for c, lab in FACTOR_LABELS.items():
        if c not in df.columns:
            continue
        s = df[c].replace(8888, pd.NA).dropna()
        # Drop any remaining out-of-range values (defensive)
        s = s[(s >= 0) & (s <= 4)]
        n = int(len(s))
        if n == 0:
            continue
        scaled = s - 2  # raw 0-4 → -2..+2
        very_pos = int((s == 4).sum())
        som_pos = int((s == 3).sum())
        som_neg = int((s == 1).sum())
        very_neg = int((s == 0).sum())
        pos = very_pos + som_pos
        neg = very_neg + som_neg
        out.append({
            "factor": lab,
            "n": n,
            "mean": round(float(scaled.mean()), 2),
            "pos_pct": round(100 * pos / n, 1),
            "neg_pct": round(100 * neg / n, 1),
        })
    out.sort(key=lambda x: -x["mean"])
    return out


def _shade(cell, fill: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _build_doc_helpers(doc: Document):
    """Return (add_heading, add_para, make_table) bound to this doc."""

    def add_heading(text, level=1, color=NAVY):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = color
        if level == 1:
            run.font.size = Pt(18)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
        elif level == 2:
            run.font.size = Pt(13)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
        else:
            run.font.size = Pt(11)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(2)
        return p

    def add_para(text, bold=False, italic=False, size=11, color=None, align=None):
        p = doc.add_paragraph()
        if align is not None:
            p.alignment = align
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        p.paragraph_format.space_after = Pt(3)
        return p

    def make_table(headers, rows, col_widths=None, header_fill="1F4E79",
                   shade_col_fills: dict | None = None):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ""
            p = hdr[i].paragraphs[0]
            run = p.add_run(h)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            _shade(hdr[i], header_fill)
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = t.rows[r_i + 1].cells[c_i]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(str(val))
                run.font.size = Pt(10)
                if shade_col_fills and c_i in shade_col_fills:
                    fill = shade_col_fills[c_i](val) if callable(shade_col_fills[c_i]) else shade_col_fills[c_i]
                    if fill:
                        _shade(cell, fill)
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in t.rows:
                    row.cells[i].width = w
        return t

    return add_heading, add_para, make_table


# ---------------------------------------------------------------------------
# Per-division report
# ---------------------------------------------------------------------------

def build_division_report(
    results: dict,
    raw_df: pd.DataFrame,
    division_name: str,
    year: Union[int, str],
    output_path: Union[str, io.BytesIO],
) -> None:
    """Generate a focused Word report for a single division.

    `division_name` matches the display label (DIVISION_MAP value), e.g. "MSA",
    "Critical Care", etc. `raw_df` is the raw REDCap DataFrame used for the
    department-level computation; this function filters it down to the division.
    """
    doc = Document()

    # Margins
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    # Default style
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    add_heading, add_para, make_table = _build_doc_helpers(doc)

    # Pull division-specific slices
    df_div = raw_df.copy()
    df_div["division_label"] = df_div["division"].map(DIVISION_MAP)
    df_div = df_div[df_div["division_label"] == division_name]

    mz_dept = results["miniz"]["clinical"]
    wb_dept = results["wellbeing"]["overall"]
    nps_dept = results["nps"]["overall"]
    ret_dept = results["retention"]["overall"]

    mz_div = next((d for d in results["miniz"]["by_division"]
                   if d["division"] == division_name), None)
    wb_div = next((d for d in results["wellbeing"]["by_division"]
                   if d["division"] == division_name), None)
    nps_div = next((d for d in results["nps"]["by_division"]
                    if d["division"] == division_name), None)
    ret_div = next((d for d in results["retention"]["by_division"]
                    if d["division"] == division_name), None)
    ld_div = next((d for d in results["leadership"]["by_division"]
                   if d["division"] == division_name), None)

    # ---------- Cover ----------
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"VUMC Department of Anesthesiology")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = TEAL

    p = doc.add_paragraph()
    run = p.add_run(f"{division_name} Division — Faculty Survey {year}")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY

    add_para(
        f"Engagement & Well-being | n = {len(df_div)} faculty respondents "
        f"({mz_div['n'] if mz_div else 0} contributed at least one MINI-Z item)",
        italic=True, size=11, color=GRAY,
    )

    # Compact narrative
    if len(df_div) == 0:
        add_para(
            f"No respondents from {division_name} in the {year} survey export. "
            f"Reach out to the survey administrator if responses are expected.",
            italic=True, size=11, color=ACCENT,
        )
        if isinstance(output_path, (str, os.PathLike)):
            doc.save(output_path)
        else:
            doc.save(output_path)
        return

    # ---------- 1. Snapshot vs Department ----------
    add_heading("Snapshot — How this division compares", 1)
    add_para(
        "Each row shows this division's value, the corresponding department "
        "value, and the difference. Green = better than department, red = "
        "worse than department, by ≥1 unit (MINI-Z scale) or ≥5 percentage "
        "points.",
        italic=True, size=10, color=GRAY,
    )

    def _color_compare(div_val, dept_val, *, lower_is_better=False,
                       unit_threshold=1.0, pp_threshold=5.0):
        if div_val is None or dept_val is None:
            return None
        delta = div_val - dept_val
        if lower_is_better:
            delta = -delta
        if delta >= max(unit_threshold, pp_threshold * 0.01 * max(1, dept_val)):
            return GOOD_FILL
        if delta <= -max(unit_threshold, pp_threshold * 0.01 * max(1, dept_val)):
            return BAD_FILL
        return WARN_FILL

    rows = []
    fills = {}
    mz_div_total = (mz_div or {}).get("total", {}) or {}
    mz_div_s1 = (mz_div or {}).get("subscale1", {}) or {}
    mz_div_s2 = (mz_div or {}).get("subscale2", {}) or {}
    rows.append([
        "MINI-Z total (clinical)",
        f"{mz_div_total.get('mean', '—'):.2f}" if mz_div_total.get('mean') is not None else "—",
        f"{mz_dept['total']['mean']:.2f}",
        _fmt_delta((mz_div_total.get('mean') or 0) - mz_dept['total']['mean']) if mz_div_total.get('mean') is not None else "—",
    ])
    rows.append([
        "Subscale 1 (supportive work env)",
        f"{mz_div_s1.get('mean', '—'):.2f}" if mz_div_s1.get('mean') is not None else "—",
        f"{mz_dept['subscale1']['mean']:.2f}",
        _fmt_delta((mz_div_s1.get('mean') or 0) - mz_dept['subscale1']['mean']) if mz_div_s1.get('mean') is not None else "—",
    ])
    rows.append([
        "Subscale 2 (work pace / EMR)",
        f"{mz_div_s2.get('mean', '—'):.2f}" if mz_div_s2.get('mean') is not None else "—",
        f"{mz_dept['subscale2']['mean']:.2f}",
        _fmt_delta((mz_div_s2.get('mean') or 0) - mz_dept['subscale2']['mean']) if mz_div_s2.get('mean') is not None else "—",
    ])
    rows.append([
        "% Burnout (Q2 ≤ 3)",
        f"{(mz_div or {}).get('burnout_pct', '—')}%" if (mz_div or {}).get('burnout_pct') is not None else "—",
        f"{mz_dept['burnout_pct']:.1f}%",
        _fmt_pct_delta(((mz_div or {}).get('burnout_pct') or 0) - mz_dept['burnout_pct']) if (mz_div or {}).get('burnout_pct') is not None else "—",
    ])
    rows.append([
        "% High stress (Q5)",
        f"{(mz_div or {}).get('stress_pct', '—')}%" if (mz_div or {}).get('stress_pct') is not None else "—",
        f"{mz_dept['stress_pct']:.1f}%",
        _fmt_pct_delta(((mz_div or {}).get('stress_pct') or 0) - mz_dept['stress_pct']) if (mz_div or {}).get('stress_pct') is not None else "—",
    ])
    rows.append([
        "% Job satisfaction (Q1)",
        f"{(mz_div or {}).get('jobsat_pct', '—')}%" if (mz_div or {}).get('jobsat_pct') is not None else "—",
        f"{mz_dept['jobsat_pct']:.1f}%",
        _fmt_pct_delta(((mz_div or {}).get('jobsat_pct') or 0) - mz_dept['jobsat_pct']) if (mz_div or {}).get('jobsat_pct') is not None else "—",
    ])
    rows.append([
        "Well-being index (mean, 0–10)",
        f"{wb_div['mean']:.2f}" if wb_div else "—",
        f"{wb_dept['mean']:.2f}",
        _fmt_delta((wb_div['mean'] if wb_div else 0) - wb_dept['mean']) if wb_div else "—",
    ])
    rows.append([
        "NPS",
        f"{nps_div['nps']:.1f}" if nps_div else "—",
        f"{nps_dept['nps']:.1f}",
        _fmt_delta((nps_div['nps'] if nps_div else 0) - nps_dept['nps']) if nps_div else "—",
    ])
    rows.append([
        "% Retention at-risk (3-year horizon)",
        f"{ret_div['at_risk_pct']:.1f}%" if ret_div else "—",
        f"{ret_dept['at_risk_pct']:.1f}%",
        _fmt_pct_delta((ret_div['at_risk_pct'] if ret_div else 0) - ret_dept['at_risk_pct']) if ret_div else "—",
    ])

    # Color the delta column
    def _delta_fill(text):
        if not isinstance(text, str) or text == "—":
            return None
        try:
            val = float(text.replace("pp", "").replace("+", ""))
        except ValueError:
            return None
        # Determine which metric (find row index by content) - simple sign rule
        # For positive metrics (MINI-Z, S1, S2, jobsat, WBI, NPS): positive=good
        # For negative metrics (burnout, stress, retention at-risk): positive=bad
        # We'll just shade by sign; manual interpretation handles the direction
        if "pp" in text:
            return GOOD_FILL if "-" in text else BAD_FILL  # negative pp = improvement for burnout/stress/risk
        return GOOD_FILL if not text.startswith("-") else BAD_FILL

    make_table(
        ["Indicator", f"{division_name}", "Department", "Δ vs Dept"],
        rows,
        col_widths=[Inches(2.9), Inches(1.2), Inches(1.2), Inches(1.2)],
    )

    add_para(
        "Note: 'Δ vs Dept' shading is informational. For burnout, high stress, "
        "and retention at-risk, lower is better — a negative delta is favorable. "
        "For MINI-Z, WBI, NPS, and job satisfaction, higher is better.",
        italic=True, size=9, color=GRAY,
    )

    doc.add_paragraph()

    # ---------- 2. Top job factors (this division) ----------
    add_heading("What's driving satisfaction in this division", 1)
    div_factors = _compute_division_factors(df_div)
    if not div_factors:
        add_para("No job-factor data available for this division.",
                 italic=True, color=GRAY)
    else:
        add_heading("Top 5 most-positive factors", 2)
        rows = [[f["factor"], f["n"], f"{f['mean']:+.2f}", f"{f['pos_pct']:.0f}%", f"{f['neg_pct']:.0f}%"]
                for f in div_factors[:5]]
        make_table(
            ["Factor", "n", "Mean (−2…+2)", "% positive", "% negative"],
            rows,
            col_widths=[Inches(3.4), Inches(0.5), Inches(1.0), Inches(1.0), Inches(1.0)],
        )
        doc.add_paragraph()

        add_heading("Top 5 most-negative factors", 2)
        rows = [[f["factor"], f["n"], f"{f['mean']:+.2f}", f"{f['pos_pct']:.0f}%", f"{f['neg_pct']:.0f}%"]
                for f in div_factors[-5:][::-1]]
        make_table(
            ["Factor", "n", "Mean (−2…+2)", "% positive", "% negative"],
            rows,
            col_widths=[Inches(3.4), Inches(0.5), Inches(1.0), Inches(1.0), Inches(1.0)],
        )
        doc.add_paragraph()

    # ---------- 3. Leadership ratings (division-specific) ----------
    add_heading("Leadership ratings — how this division sees its leaders", 1)
    if not ld_div:
        add_para("No leadership ratings available for this division.",
                 italic=True, color=GRAY)
    else:
        rows = []
        for col, item_label, scale in LEADERSHIP_ITEMS:
            mean_v = ld_div.get(f"{col}_mean")
            n_v = ld_div.get(f"{col}_n")
            agree_v = ld_div.get(f"{col}_agree_pct")
            if mean_v is None or n_v is None or n_v == 0:
                continue
            short_label = LEADERSHIP_LABELS.get(col, item_label)
            rows.append([
                short_label,
                n_v,
                f"{mean_v:.2f}",
                f"{agree_v:.1f}%" if agree_v is not None else "—",
            ])
        if rows:
            make_table(
                ["Leadership item", "n", "Mean (1–5)", "% Agree+"],
                rows,
                col_widths=[Inches(4.0), Inches(0.5), Inches(1.0), Inches(1.0)],
            )
        else:
            add_para("Leadership items not answered by this division's "
                     "respondents.", italic=True, color=GRAY)
    doc.add_paragraph()

    # ---------- 4. Discussion prompts ----------
    add_heading("Suggested discussion prompts (for the Division Chief)", 1)
    prompts = [
        f"The two indicators where {division_name} differs most from the "
        "department average — what's driving the gap, and what's within the "
        "division's control?",
        "Are the most-positive factors in this division being protected? "
        "What recent decisions might have reinforced (or risked) them?",
        "Among the most-negative factors, which one or two are realistically "
        "addressable within the next year by this division?",
        "Are the leadership ratings for the division chief role consistent "
        "with what you'd expect? If lower than the department-wide average, "
        "what's the underlying signal — communication, decision-making, "
        "presence, advocacy?",
        "For division members at retention risk: what specific changes would "
        "most influence their decision to stay? (The department-wide stay "
        "drivers were compensation, length of clinical days, schedule "
        "predictability, and academic-time receipt.)",
    ]
    for prompt in prompts:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(prompt).font.size = Pt(11)

    # ---------- Footer ----------
    add_para(
        f"Generated from REDCap export — {year} survey cycle. "
        "REDCap responses are anonymous; small-cell divisions (n < 10) "
        "may not be statistically robust and should be interpreted alongside "
        "qualitative input from the division chief.",
        italic=True, size=9, color=GRAY,
    )

    if isinstance(output_path, (str, os.PathLike)):
        doc.save(output_path)
    else:
        doc.save(output_path)


# ---------------------------------------------------------------------------
# All-division ZIP
# ---------------------------------------------------------------------------

def _safe_filename(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9._-]+", "_", name).strip("_")


def build_all_division_reports(
    results: dict,
    raw_df: pd.DataFrame,
    year: Union[int, str],
    divisions: Iterable[str] | None = None,
) -> bytes:
    """Build a Word doc per division and return a ZIP of all of them as bytes.

    `divisions` defaults to the division list found in
    `results["miniz"]["by_division"]`. Pass an explicit iterable to override.
    """
    if divisions is None:
        divisions = [d["division"] for d in results["miniz"]["by_division"]
                     if d["division"] and d["division"] != "Missing"]

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for div in divisions:
            doc_buffer = io.BytesIO()
            build_division_report(results, raw_df, div, year, doc_buffer)
            doc_buffer.seek(0)
            safe = _safe_filename(str(div))
            zf.writestr(
                f"VUMC_Anesth_{year}_{safe}_Division_Report.docx",
                doc_buffer.getvalue(),
            )
    zip_buffer.seek(0)
    return zip_buffer.getvalue()
