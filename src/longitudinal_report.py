"""Generate a streamlined longitudinal Word report covering year-over-year
changes across all metrics — items, subscales, WBI, NPS, retention, job
factors, and leadership.

Charts are rendered to PNGs into a tempdir and embedded in the .docx.
"""
from __future__ import annotations

import os
import tempfile
from io import BytesIO
from typing import Dict, Iterable, Union

import matplotlib
matplotlib.use("Agg")  # non-interactive backend for server contexts
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# Colors
NAVY = "#1F4E79"
TEAL = "#2E8B8B"
ACCENT = "#C0504D"
GOLD = "#C9A227"
PURPLE = "#6A4C93"
GRAY = "#7F7F7F"
PALETTE = [NAVY, TEAL, GOLD, ACCENT, PURPLE]

# python-docx colors
NAVY_RGB = RGBColor(0x1F, 0x4E, 0x79)
TEAL_RGB = RGBColor(0x2E, 0x8B, 0x8B)
GRAY_RGB = RGBColor(0x55, 0x55, 0x55)
ACCENT_RGB = RGBColor(0xC0, 0x50, 0x4D)


# ---------------------------------------------------------------------------
# Generic helpers
# ---------------------------------------------------------------------------

def _sort_key(y):
    try:
        return (0, int(y))
    except Exception:
        return (1, str(y))


def _fmt_delta(d, *, pp=False, decimals=2):
    if d is None or (isinstance(d, float) and np.isnan(d)):
        return "—"
    sign = "+" if d >= 0 else ""
    return f"{sign}{d:.1f}pp" if pp else f"{sign}{d:.{decimals}f}"


def _safe(val):
    return "—" if val is None else val


# ---------------------------------------------------------------------------
# Chart helpers (matplotlib PNGs)
# ---------------------------------------------------------------------------

def _save_fig(fig, path):
    fig.tight_layout()
    fig.savefig(path, dpi=140, bbox_inches="tight", facecolor="white")
    plt.close(fig)


def _chart_headline_metrics(years, miniz_total, burnout_pct, wbi_mean, nps,
                            path):
    fig, axes = plt.subplots(2, 2, figsize=(10, 6))
    specs = [
        (axes[0, 0], "MINI-Z Total (Clinical)", "Mean (10–50)",
         miniz_total, NAVY, [10, 50], 40, "Joyful (≥40)"),
        (axes[0, 1], "% Clinical Burnout", "% (Q2 ≤ 3)",
         burnout_pct, ACCENT, [0, 100], None, None),
        (axes[1, 0], "Well-being Index (Overall)", "Mean (0–10)",
         wbi_mean, TEAL, [0, 10], 8, "High (≥8)"),
        (axes[1, 1], "Net Promoter Score (Overall)", "NPS",
         nps, GOLD, [-100, 100], 0, None),
    ]
    for ax, title, ylab, vals, color, ylim, threshold, thr_label in specs:
        ax.plot(years, vals, color=color, marker="o", markersize=8, lw=2.5)
        for x, v in zip(years, vals):
            if v is not None:
                ax.annotate(f"{v:.1f}", (x, v), textcoords="offset points",
                            xytext=(0, 8), ha="center", fontsize=9,
                            fontweight="bold")
        if threshold is not None:
            ax.axhline(threshold, ls="--", color=ACCENT, lw=1, alpha=0.7)
            if thr_label:
                ax.text(years[-1], threshold, f" {thr_label}",
                        va="center", fontsize=8, color=ACCENT)
        ax.set_title(title, fontweight="bold", color=NAVY, fontsize=11)
        ax.set_xlabel("Year")
        ax.set_ylabel(ylab)
        ax.set_ylim(ylim)
        ax.grid(axis="y", linestyle=":", alpha=0.5)
        ax.spines["top"].set_visible(False)
        ax.spines["right"].set_visible(False)
    _save_fig(fig, path)


def _chart_subscale_trends(years, s1_vals, s2_vals, path):
    fig, ax = plt.subplots(figsize=(8, 4.6))
    ax.plot(years, s1_vals, color=NAVY, marker="o", markersize=9, lw=2.5,
            label="Subscale 1 — Supportive Work Env")
    ax.plot(years, s2_vals, color=TEAL, marker="s", markersize=9, lw=2.5,
            label="Subscale 2 — Work Pace / EMR Stress")
    for x, v in zip(years, s1_vals):
        if v is not None:
            ax.annotate(f"{v:.2f}", (x, v), xytext=(0, 8),
                        textcoords="offset points", ha="center",
                        fontsize=9, fontweight="bold", color=NAVY)
    for x, v in zip(years, s2_vals):
        if v is not None:
            ax.annotate(f"{v:.2f}", (x, v), xytext=(0, -14),
                        textcoords="offset points", ha="center",
                        fontsize=9, fontweight="bold", color=TEAL)
    ax.axhline(20, ls="--", color=ACCENT, lw=1, alpha=0.7)
    ax.text(years[-1], 20.2, " ≥20 highly supportive / reasonable pace",
            color=ACCENT, fontsize=8, va="bottom")
    ax.set_title("MINI-Z Subscale Means by Year — Clinical Faculty",
                 fontweight="bold", color=NAVY, fontsize=12)
    ax.set_xlabel("Year")
    ax.set_ylabel("Subscale mean (5–25)")
    ax.set_ylim(5, 25)
    ax.legend(loc="lower right", fontsize=9, frameon=False)
    ax.grid(axis="y", linestyle=":", alpha=0.5)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    _save_fig(fig, path)


def _chart_miniz_items(years, items_means, path):
    """items_means: list of (item_short_label, [vals_per_year])."""
    short_labels = [a for a, _ in items_means]
    n_items = len(short_labels)
    n_years = len(years)
    x = np.arange(n_items)
    width = 0.8 / max(n_years, 1)

    fig, ax = plt.subplots(figsize=(10, 4.6))
    for i, y in enumerate(years):
        vals = [vals_per_year[i] for _, vals_per_year in items_means]
        offset = (i - (n_years - 1) / 2) * width
        bars = ax.bar(x + offset, vals, width,
                      color=PALETTE[i % len(PALETTE)], label=str(y))
        for b, v in zip(bars, vals):
            if v is not None:
                ax.annotate(f"{v:.2f}", (b.get_x() + b.get_width() / 2,
                                          b.get_height()),
                            xytext=(0, 3), textcoords="offset points",
                            ha="center", fontsize=7.5)
    ax.set_xticks(x)
    ax.set_xticklabels(short_labels, fontsize=9)
    ax.set_title("MINI-Z Item Means by Year (overall sample)",
                 fontweight="bold", color=NAVY, fontsize=12)
    ax.set_xlabel("MINI-Z item")
    ax.set_ylabel("Mean (1–5; 5 = best)")
    ax.set_ylim(0, 5.5)
    ax.legend(loc="upper center", ncol=n_years,
              bbox_to_anchor=(0.5, -0.12), frameon=False, fontsize=9)
    ax.grid(axis="y", linestyle=":", alpha=0.5)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    _save_fig(fig, path)


def _chart_wbi_bands(years, low_pcts, mid_pcts, high_pcts, path):
    fig, ax = plt.subplots(figsize=(7, 4.3))
    x = np.arange(len(years))
    ax.bar(x, low_pcts, color=ACCENT, label="Low (0–4)")
    ax.bar(x, mid_pcts, bottom=low_pcts, color=GOLD, label="Mid (5–7)")
    ax.bar(x, high_pcts,
           bottom=[a + b for a, b in zip(low_pcts, mid_pcts)],
           color=TEAL, label="High (8–10)")
    for i, (lo, mi, hi) in enumerate(zip(low_pcts, mid_pcts, high_pcts)):
        if lo > 4:
            ax.text(i, lo / 2, f"{lo:.0f}%", ha="center", va="center",
                    color="white", fontweight="bold", fontsize=9)
        if mi > 4:
            ax.text(i, lo + mi / 2, f"{mi:.0f}%", ha="center", va="center",
                    color="white", fontweight="bold", fontsize=9)
        if hi > 4:
            ax.text(i, lo + mi + hi / 2, f"{hi:.0f}%", ha="center",
                    va="center", color="white", fontweight="bold",
                    fontsize=9)
    ax.set_xticks(x)
    ax.set_xticklabels([str(y) for y in years])
    ax.set_title("Well-being Index — Band Distribution by Year",
                 fontweight="bold", color=NAVY, fontsize=12)
    ax.set_xlabel("Year")
    ax.set_ylabel("% of respondents")
    ax.set_ylim(0, 105)
    ax.legend(loc="upper center", ncol=3, bbox_to_anchor=(0.5, -0.12),
              frameon=False, fontsize=9)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    _save_fig(fig, path)


def _chart_nps_components(years, prom_pcts, pass_pcts, det_pcts,
                          nps_vals, path):
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4.4),
                                   gridspec_kw={"width_ratios": [1, 1]})
    x = np.arange(len(years))
    ax1.bar(x, det_pcts, color=ACCENT, label="Detractors (0–6)")
    ax1.bar(x, pass_pcts, bottom=det_pcts, color=GOLD,
            label="Passives (7–8)")
    ax1.bar(x, prom_pcts,
            bottom=[a + b for a, b in zip(det_pcts, pass_pcts)],
            color=TEAL, label="Promoters (9–10)")
    for i, (d, p, pr) in enumerate(zip(det_pcts, pass_pcts, prom_pcts)):
        if d > 4:
            ax1.text(i, d / 2, f"{d:.0f}%", ha="center", va="center",
                     color="white", fontweight="bold", fontsize=9)
        if p > 4:
            ax1.text(i, d + p / 2, f"{p:.0f}%", ha="center", va="center",
                     color="white", fontweight="bold", fontsize=9)
        if pr > 4:
            ax1.text(i, d + p + pr / 2, f"{pr:.0f}%", ha="center",
                     va="center", color="white", fontweight="bold",
                     fontsize=9)
    ax1.set_xticks(x)
    ax1.set_xticklabels([str(y) for y in years])
    ax1.set_title("NPS Component Distribution",
                  fontweight="bold", color=NAVY, fontsize=11)
    ax1.set_ylabel("% of respondents")
    ax1.set_ylim(0, 105)
    ax1.legend(loc="upper center", ncol=3, bbox_to_anchor=(0.5, -0.12),
               frameon=False, fontsize=8)
    ax1.spines["top"].set_visible(False)
    ax1.spines["right"].set_visible(False)

    ax2.plot(years, nps_vals, color=GOLD, marker="o", markersize=10, lw=3)
    for x_val, v in zip(years, nps_vals):
        if v is not None:
            ax2.annotate(f"{v:.1f}", (x_val, v), xytext=(0, 10),
                         textcoords="offset points", ha="center",
                         fontsize=10, fontweight="bold")
    ax2.axhline(0, color="black", lw=1)
    ax2.set_title("NPS Trend", fontweight="bold", color=NAVY, fontsize=11)
    ax2.set_xlabel("Year")
    ax2.set_ylabel("NPS")
    ax2.set_ylim(-100, 100)
    ax2.grid(axis="y", linestyle=":", alpha=0.5)
    ax2.spines["top"].set_visible(False)
    ax2.spines["right"].set_visible(False)
    _save_fig(fig, path)


def _chart_retention(years, at_risk_pcts, likely_stay_pcts, means, path):
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4.4))
    x = np.arange(len(years))
    width = 0.36
    bars1 = ax1.bar(x - width / 2, at_risk_pcts, width, color=ACCENT,
                    label="% at-risk")
    bars2 = ax1.bar(x + width / 2, likely_stay_pcts, width, color=NAVY,
                    label="% likely-to-stay")
    for b, v in zip(bars1, at_risk_pcts):
        if v is not None:
            ax1.annotate(f"{v:.0f}%", (b.get_x() + b.get_width() / 2,
                                        b.get_height()),
                         xytext=(0, 3), textcoords="offset points",
                         ha="center", fontsize=9, fontweight="bold")
    for b, v in zip(bars2, likely_stay_pcts):
        if v is not None:
            ax1.annotate(f"{v:.0f}%", (b.get_x() + b.get_width() / 2,
                                        b.get_height()),
                         xytext=(0, 3), textcoords="offset points",
                         ha="center", fontsize=9, fontweight="bold")
    ax1.set_xticks(x)
    ax1.set_xticklabels([str(y) for y in years])
    ax1.set_title("Retention — At-risk vs Likely-to-stay",
                  fontweight="bold", color=NAVY, fontsize=11)
    ax1.set_ylabel("% of respondents")
    ax1.set_ylim(0, 100)
    ax1.legend(loc="upper right", fontsize=9, frameon=False)
    ax1.spines["top"].set_visible(False)
    ax1.spines["right"].set_visible(False)

    ax2.plot(years, means, color=NAVY, marker="o", markersize=10, lw=3)
    for x_val, v in zip(years, means):
        if v is not None:
            ax2.annotate(f"{v:.2f}", (x_val, v), xytext=(0, 10),
                         textcoords="offset points", ha="center",
                         fontsize=10, fontweight="bold")
    ax2.set_title("Mean Retention Score",
                  fontweight="bold", color=NAVY, fontsize=11)
    ax2.set_xlabel("Year")
    ax2.set_ylabel("Mean (0=very unlikely to stay … 4=very likely)")
    ax2.set_ylim(0, 4)
    ax2.grid(axis="y", linestyle=":", alpha=0.5)
    ax2.spines["top"].set_visible(False)
    ax2.spines["right"].set_visible(False)
    _save_fig(fig, path)


def _chart_grouped_bars(years, labels_and_values, title, ylabel, path,
                        ylim=(-2, 2), zero_line=True, figsize=(10, 5)):
    """labels_and_values: list of (short_label, [val_per_year])."""
    short_labels = [a for a, _ in labels_and_values]
    n_items = len(short_labels)
    n_years = len(years)
    x = np.arange(n_items)
    width = 0.8 / max(n_years, 1)
    fig, ax = plt.subplots(figsize=figsize)
    for i, y in enumerate(years):
        vals = [vals_per_year[i] for _, vals_per_year in labels_and_values]
        offset = (i - (n_years - 1) / 2) * width
        bars = ax.bar(x + offset, vals, width,
                      color=PALETTE[i % len(PALETTE)], label=str(y))
        for b, v in zip(bars, vals):
            if v is not None:
                yv = b.get_height()
                offset_y = 3 if yv >= 0 else -10
                ax.annotate(f"{v:+.2f}", (b.get_x() + b.get_width() / 2, yv),
                            xytext=(0, offset_y),
                            textcoords="offset points", ha="center",
                            fontsize=7.5)
    if zero_line:
        ax.axhline(0, color="black", lw=1)
    ax.set_xticks(x)
    ax.set_xticklabels(short_labels, fontsize=8, rotation=20, ha="right")
    ax.set_title(title, fontweight="bold", color=NAVY, fontsize=12)
    ax.set_ylabel(ylabel)
    ax.set_ylim(ylim)
    ax.legend(loc="upper center", ncol=n_years,
              bbox_to_anchor=(0.5, -0.30), frameon=False, fontsize=9)
    ax.grid(axis="y", linestyle=":", alpha=0.5)
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    _save_fig(fig, path)


# ---------------------------------------------------------------------------
# Word doc helpers
# ---------------------------------------------------------------------------

def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _build_doc_helpers(doc):
    def add_heading(text, level=1, color=NAVY_RGB):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = color
        sizes = {1: 18, 2: 14, 3: 12}
        space_before = {1: 14, 2: 10, 3: 6}
        space_after = {1: 6, 2: 4, 3: 2}
        run.font.size = Pt(sizes.get(level, 11))
        p.paragraph_format.space_before = Pt(space_before.get(level, 4))
        p.paragraph_format.space_after = Pt(space_after.get(level, 2))
        return p

    def add_para(text, bold=False, italic=False, size=11, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = color
        p.paragraph_format.space_after = Pt(4)
        return p

    def make_table(headers, rows, col_widths=None, header_fill="1F4E79"):
        t = doc.add_table(rows=1 + len(rows), cols=len(headers))
        t.style = "Light Grid Accent 1"
        hdr = t.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = ""
            p = hdr[i].paragraphs[0]
            r = p.add_run(h)
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            r.font.size = Pt(10)
            _shade(hdr[i], header_fill)
        for ri, row in enumerate(rows):
            for ci, val in enumerate(row):
                cell = t.rows[ri + 1].cells[ci]
                cell.text = ""
                p = cell.paragraphs[0]
                r = p.add_run("" if val is None else str(val))
                r.font.size = Pt(10)
        if col_widths:
            for i, w in enumerate(col_widths):
                for row in t.rows:
                    row.cells[i].width = w
        return t

    def add_image(path, width_in=6.5):
        if os.path.exists(path):
            doc.add_picture(path, width=Inches(width_in))
            last_p = doc.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return add_heading, add_para, make_table, add_image


# ---------------------------------------------------------------------------
# Main report builder
# ---------------------------------------------------------------------------

def build_longitudinal_report(
    results_by_year: Dict[str, dict],
    output_path: Union[str, BytesIO],
) -> None:
    """Generate a streamlined Word report covering all metrics across years."""
    if not results_by_year or len(results_by_year) < 2:
        raise ValueError("Need at least two years of data for a longitudinal "
                         "report.")

    years = sorted(results_by_year.keys(), key=_sort_key)
    first_y, last_y = years[0], years[-1]
    years_str = [str(y) for y in years]
    delta_col = f"Δ {first_y}→{last_y}"

    # Setup tempdir for charts
    tmpdir = tempfile.mkdtemp(prefix="longreport_")

    # ----- Build all charts upfront -----
    miniz_totals = [results_by_year[y]["miniz"]["clinical"].get("total", {}).get("mean")
                    for y in years]
    burnout_pcts = [results_by_year[y]["miniz"]["clinical"].get("burnout_pct")
                    for y in years]
    wbi_means = [(results_by_year[y]["wellbeing"]["overall"] or {}).get("mean")
                 for y in years]
    nps_vals = [(results_by_year[y]["nps"]["overall"] or {}).get("nps")
                for y in years]

    p_head = os.path.join(tmpdir, "headline.png")
    _chart_headline_metrics(years_str, miniz_totals, burnout_pcts, wbi_means,
                             nps_vals, p_head)

    # Subscale chart
    s1_vals = [results_by_year[y]["miniz"]["clinical"].get("subscale1", {}).get("mean")
               for y in years]
    s2_vals = [results_by_year[y]["miniz"]["clinical"].get("subscale2", {}).get("mean")
               for y in years]
    p_subscale = os.path.join(tmpdir, "subscale.png")
    _chart_subscale_trends(years_str, s1_vals, s2_vals, p_subscale)

    # Item-level chart
    base_items = results_by_year[first_y]["miniz"]["overall"]["items"]
    item_labels = [it["item"] for it in base_items]
    short_labels = [lab.split(".")[0] if "." in lab else lab[:6]
                    for lab in item_labels]
    items_means = []
    for label, short in zip(item_labels, short_labels):
        vals = []
        for y in years:
            idx = {it["item"]: it for it in
                   results_by_year[y]["miniz"]["overall"]["items"]}
            vals.append((idx.get(label) or {}).get("mean"))
        items_means.append((short, vals))
    p_items = os.path.join(tmpdir, "miniz_items.png")
    _chart_miniz_items(years_str, items_means, p_items)

    # WBI band chart
    low_pcts = [(results_by_year[y]["wellbeing"]["overall"] or {}).get("low_pct", 0)
                for y in years]
    mid_pcts = [(results_by_year[y]["wellbeing"]["overall"] or {}).get("mid_pct", 0)
                for y in years]
    high_pcts = [(results_by_year[y]["wellbeing"]["overall"] or {}).get("high_pct", 0)
                 for y in years]
    p_wbi = os.path.join(tmpdir, "wbi_bands.png")
    _chart_wbi_bands(years_str, low_pcts, mid_pcts, high_pcts, p_wbi)

    # NPS components chart
    prom_pcts = [(results_by_year[y]["nps"]["overall"] or {}).get("promoters_pct", 0)
                 for y in years]
    pass_pcts = [(results_by_year[y]["nps"]["overall"] or {}).get("passives_pct", 0)
                 for y in years]
    det_pcts = [(results_by_year[y]["nps"]["overall"] or {}).get("detractors_pct", 0)
                for y in years]
    p_nps = os.path.join(tmpdir, "nps_comp.png")
    _chart_nps_components(years_str, prom_pcts, pass_pcts, det_pcts,
                          nps_vals, p_nps)

    # Retention chart
    at_risk_pcts = [(results_by_year[y]["retention"]["overall"] or {}).get("at_risk_pct")
                    for y in years]
    likely_stay_pcts = [(results_by_year[y]["retention"]["overall"] or {}).get("likely_stay_pct")
                        for y in years]
    ret_means = [(results_by_year[y]["retention"]["overall"] or {}).get("mean")
                 for y in years]
    p_retention = os.path.join(tmpdir, "retention.png")
    _chart_retention(years_str, at_risk_pcts, likely_stay_pcts, ret_means,
                     p_retention)

    # Factor movers chart (top 10 by absolute change)
    all_factor_labels = sorted({f["factor"] for y in years
                                for f in results_by_year[y]["factors"]
                                if f["mean"] is not None})
    factor_delta_pairs = []
    for label in all_factor_labels:
        fi = {f["factor"]: f for f in results_by_year[first_y]["factors"]}
        li = {f["factor"]: f for f in results_by_year[last_y]["factors"]}
        fm = (fi.get(label) or {}).get("mean")
        lm = (li.get(label) or {}).get("mean")
        if fm is None or lm is None:
            continue
        factor_delta_pairs.append((label, fm, lm, lm - fm))
    factor_delta_pairs.sort(key=lambda r: abs(r[3]), reverse=True)
    top_movers = factor_delta_pairs[:10]
    if top_movers:
        movers_labels = []
        for label, _, _, _ in top_movers:
            short = label if len(label) <= 28 else label[:28] + "…"
            vals = []
            for y in years:
                fi = {f["factor"]: f for f in results_by_year[y]["factors"]}
                vals.append((fi.get(label) or {}).get("mean"))
            movers_labels.append((short, vals))
        p_factor_movers = os.path.join(tmpdir, "factor_movers.png")
        _chart_grouped_bars(years_str, movers_labels,
                             title="Top 10 Job-Factor Movers",
                             ylabel="Mean impact (−2 negative → +2 positive)",
                             path=p_factor_movers, ylim=(-2, 2),
                             zero_line=True, figsize=(10, 5.5))
    else:
        p_factor_movers = None

    # Leadership movers chart (top 8 by absolute change)
    ld_items_by_year = {y: results_by_year[y]["leadership"]["overall"]
                        for y in years}
    all_ld_items = sorted({it["item"] for y in years
                           for it in ld_items_by_year[y]})
    ld_delta_pairs = []
    for label in all_ld_items:
        fi = {it["item"]: it for it in ld_items_by_year[first_y]}
        li = {it["item"]: it for it in ld_items_by_year[last_y]}
        fm = (fi.get(label) or {}).get("mean")
        lm = (li.get(label) or {}).get("mean")
        if fm is None or lm is None:
            continue
        ld_delta_pairs.append((label, fm, lm, lm - fm))
    ld_delta_pairs.sort(key=lambda r: abs(r[3]), reverse=True)
    top_ld_movers = ld_delta_pairs[:8]
    if top_ld_movers:
        movers_labels = []
        for label, _, _, _ in top_ld_movers:
            short = label if len(label) <= 30 else label[:30] + "…"
            vals = []
            for y in years:
                ld_idx = {it["item"]: it for it in ld_items_by_year[y]}
                vals.append((ld_idx.get(label) or {}).get("mean"))
            movers_labels.append((short, vals))
        p_ld_movers = os.path.join(tmpdir, "ld_movers.png")
        _chart_grouped_bars(years_str, movers_labels,
                             title="Top 8 Leadership-Item Movers",
                             ylabel="Mean rating (1–5)",
                             path=p_ld_movers, ylim=(1, 5),
                             zero_line=False, figsize=(10, 5))
    else:
        p_ld_movers = None

    # ----- Build the Word doc -----
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    add_heading, add_para, make_table, add_image = _build_doc_helpers(doc)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("VUMC Department of Anesthesiology")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = TEAL_RGB
    p = doc.add_paragraph()
    run = p.add_run(f"Faculty Survey — Longitudinal Comparison "
                    f"({first_y} → {last_y})")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = NAVY_RGB
    sample_info = " · ".join(
        f"{y} n = {results_by_year[y]['sample']['n_total']}" for y in years
    )
    add_para(f"Engagement & Well-being | {sample_info}",
             italic=True, size=11, color=GRAY_RGB)

    # ----- 1. Executive summary -----
    add_heading("Executive Summary", 1)
    bullets = []

    def _arrow(d, *, lower_is_better=False):
        if d is None or (isinstance(d, float) and np.isnan(d)):
            return ""
        good = (d < 0) if lower_is_better else (d > 0)
        return " ▲" if good else (" ▼" if d != 0 else " →")

    d_mz = (miniz_totals[-1] - miniz_totals[0]) if (miniz_totals[0] is not None
                                                     and miniz_totals[-1] is not None) else None
    d_bo = (burnout_pcts[-1] - burnout_pcts[0]) if (burnout_pcts[0] is not None
                                                    and burnout_pcts[-1] is not None) else None
    d_wb = (wbi_means[-1] - wbi_means[0]) if (wbi_means[0] is not None
                                              and wbi_means[-1] is not None) else None
    d_nps = (nps_vals[-1] - nps_vals[0]) if (nps_vals[0] is not None
                                              and nps_vals[-1] is not None) else None
    d_risk = (at_risk_pcts[-1] - at_risk_pcts[0]) if (at_risk_pcts[0] is not None
                                                       and at_risk_pcts[-1] is not None) else None

    bullets.append(
        f"MINI-Z total (clinical): {miniz_totals[0]:.2f} → {miniz_totals[-1]:.2f} "
        f"({_fmt_delta(d_mz)}{_arrow(d_mz)}). Threshold for joyful workplace is 40."
    )
    bullets.append(
        f"% Burnout (clinical): {burnout_pcts[0]:.1f}% → {burnout_pcts[-1]:.1f}% "
        f"({_fmt_delta(d_bo, pp=True)}{_arrow(d_bo, lower_is_better=True)})."
    )
    bullets.append(
        f"Well-being Index (overall mean, 0–10): {wbi_means[0]:.2f} → {wbi_means[-1]:.2f} "
        f"({_fmt_delta(d_wb)}{_arrow(d_wb)})."
    )
    bullets.append(
        f"Net Promoter Score (overall): {nps_vals[0]:.1f} → {nps_vals[-1]:.1f} "
        f"({_fmt_delta(d_nps)}{_arrow(d_nps)})."
    )
    if d_risk is not None:
        bullets.append(
            f"Retention at-risk (3-year horizon): {at_risk_pcts[0]:.1f}% → "
            f"{at_risk_pcts[-1]:.1f}% "
            f"({_fmt_delta(d_risk, pp=True)}{_arrow(d_risk, lower_is_better=True)})."
        )
    for b in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(b).font.size = Pt(11)

    doc.add_paragraph()

    # ----- 2. Headline metrics -----
    add_heading("1. Headline Metrics", 1)
    add_image(p_head, width_in=6.5)
    add_para("Figure 1. Headline indicators across loaded survey years.",
             italic=True, size=10, color=GRAY_RGB)

    rows = []
    for y in years:
        rows.append([
            str(y),
            results_by_year[y]["sample"]["n_total"],
            f"{miniz_totals[years.index(y)]:.2f}" if miniz_totals[years.index(y)] is not None else "—",
            f"{burnout_pcts[years.index(y)]:.1f}%" if burnout_pcts[years.index(y)] is not None else "—",
            f"{wbi_means[years.index(y)]:.2f}" if wbi_means[years.index(y)] is not None else "—",
            f"{nps_vals[years.index(y)]:.1f}" if nps_vals[years.index(y)] is not None else "—",
        ])
    make_table(["Year", "n", "MINI-Z total (clin)", "% burnout (clin)",
                "WBI mean", "NPS"], rows,
               col_widths=[Inches(0.7), Inches(0.6), Inches(1.3),
                           Inches(1.3), Inches(1.0), Inches(0.8)])
    doc.add_paragraph()

    # ----- 3. MINI-Z subscales -----
    add_heading("2. MINI-Z Subscale Trends — Clinical Faculty", 1)
    add_image(p_subscale, width_in=6.5)
    add_para("Figure 2. Subscale 1 (Supportive Work Env, Q1+Q2+Q3+Q4+Q9) and "
             "Subscale 2 (Work Pace/EMR Stress, Q5+Q6+Q7+Q8+Q10) per the "
             "Mini-Z 2.0 manual.",
             italic=True, size=10, color=GRAY_RGB)
    sub_rows = []
    for y in years:
        mz = results_by_year[y]["miniz"]["clinical"]
        s1 = mz.get("subscale1") or {}
        s2 = mz.get("subscale2") or {}
        tot = mz.get("total") or {}
        sub_rows.append([
            str(y),
            tot.get("n", "—"),
            f"{tot.get('mean', 0):.2f}" if tot.get("mean") is not None else "—",
            f"{tot.get('joyful_pct', 0):.1f}%" if tot.get("joyful_pct") is not None else "—",
            f"{s1.get('mean', 0):.2f}" if s1.get("mean") is not None else "—",
            f"{s1.get('supportive_pct', 0):.1f}%" if s1.get("supportive_pct") is not None else "—",
            f"{s2.get('mean', 0):.2f}" if s2.get("mean") is not None else "—",
            f"{s2.get('lowstress_pct', 0):.1f}%" if s2.get("lowstress_pct") is not None else "—",
        ])
    make_table(["Year", "n", "Total (10–50)", "% ≥40 joyful",
                "S1 (5–25)", "S1 % ≥20", "S2 (5–25)", "S2 % ≥20"],
               sub_rows,
               col_widths=[Inches(0.6), Inches(0.5), Inches(0.9),
                           Inches(0.8), Inches(0.7), Inches(0.7),
                           Inches(0.7), Inches(0.7)])
    doc.add_paragraph()

    # ----- 4. MINI-Z item-level -----
    add_heading("3. MINI-Z Item-Level Changes", 1)
    add_image(p_items, width_in=6.5)
    add_para("Figure 3. Item-level means across years (overall sample). "
             "Higher = better on every item.",
             italic=True, size=10, color=GRAY_RGB)
    item_rows = []
    first_idx = {it["item"]: it for it in
                 results_by_year[first_y]["miniz"]["overall"]["items"]}
    last_idx = {it["item"]: it for it in
                results_by_year[last_y]["miniz"]["overall"]["items"]}
    for label in item_labels:
        row = [label]
        for y in years:
            idx = {it["item"]: it for it in
                   results_by_year[y]["miniz"]["overall"]["items"]}
            it = idx.get(label) or {}
            m = it.get("mean")
            row.append(f"{m:.2f}" if m is not None else "—")
        fm = (first_idx.get(label) or {}).get("mean")
        lm = (last_idx.get(label) or {}).get("mean")
        row.append(_fmt_delta(lm - fm) if (fm is not None and lm is not None) else "—")
        item_rows.append(row)
    headers = ["MINI-Z item"] + [str(y) for y in years] + [delta_col]
    widths = [Inches(3.4)] + [Inches(0.7)] * len(years) + [Inches(0.8)]
    make_table(headers, item_rows, col_widths=widths)
    doc.add_paragraph()

    # ----- 5. Well-being index -----
    add_heading("4. Well-being Index", 1)
    add_image(p_wbi, width_in=6.0)
    add_para("Figure 4. Distribution of WBI scores across bands (low 0–4, "
             "mid 5–7, high 8–10) by year.",
             italic=True, size=10, color=GRAY_RGB)
    wbi_rows = []
    for y in years:
        wb = results_by_year[y]["wellbeing"]["overall"] or {}
        wbi_rows.append([
            str(y),
            wb.get("n", "—"),
            f"{wb.get('mean', 0):.2f}" if wb.get("mean") is not None else "—",
            f"{wb.get('low_pct', 0):.1f}%" if wb.get("low_pct") is not None else "—",
            f"{wb.get('mid_pct', 0):.1f}%" if wb.get("mid_pct") is not None else "—",
            f"{wb.get('high_pct', 0):.1f}%" if wb.get("high_pct") is not None else "—",
        ])
    make_table(["Year", "n", "Mean", "% Low (0–4)", "% Mid (5–7)", "% High (8–10)"],
               wbi_rows,
               col_widths=[Inches(0.7), Inches(0.5), Inches(0.7),
                           Inches(1.1), Inches(1.1), Inches(1.2)])
    doc.add_paragraph()

    # ----- 6. NPS -----
    add_heading("5. Net Promoter Score", 1)
    add_image(p_nps, width_in=6.5)
    add_para("Figure 5. NPS components (Detractors 0–6, Passives 7–8, "
             "Promoters 9–10) and overall NPS trend.",
             italic=True, size=10, color=GRAY_RGB)
    nps_rows = []
    for y in years:
        nps_d = results_by_year[y]["nps"]["overall"] or {}
        nps_rows.append([
            str(y),
            nps_d.get("n", "—"),
            f"{nps_d.get('promoters_pct', 0):.1f}%",
            f"{nps_d.get('passives_pct', 0):.1f}%",
            f"{nps_d.get('detractors_pct', 0):.1f}%",
            f"{nps_d.get('nps', 0):.1f}",
        ])
    make_table(["Year", "n", "% Promoters", "% Passives", "% Detractors", "NPS"],
               nps_rows,
               col_widths=[Inches(0.7), Inches(0.5), Inches(1.1),
                           Inches(1.0), Inches(1.1), Inches(0.7)])
    doc.add_paragraph()

    # ----- 7. Retention -----
    add_heading("6. Retention Risk (3-year horizon)", 1)
    add_image(p_retention, width_in=6.5)
    add_para("Figure 6. Retention indicators across years. Item: "
             "\"In 3 years, I will still be working at VUMC.\"",
             italic=True, size=10, color=GRAY_RGB)
    ret_rows = []
    for y in years:
        ret = results_by_year[y]["retention"]["overall"] or {}
        ret_rows.append([
            str(y),
            ret.get("n", "—"),
            f"{ret.get('mean', 0):.2f}" if ret.get("mean") is not None else "—",
            f"{ret.get('at_risk_pct', 0):.1f}%",
            f"{ret.get('likely_stay_pct', 0):.1f}%",
        ])
    make_table(["Year", "n", "Mean (0–4)", "% At-risk", "% Likely-to-stay"],
               ret_rows,
               col_widths=[Inches(0.7), Inches(0.5), Inches(0.9),
                           Inches(1.1), Inches(1.4)])
    doc.add_paragraph()

    # ----- 8. Job factors -----
    add_heading("7. Job-Factor Changes", 1)
    add_para(
        f"Each factor was rated for impact on satisfaction on a 5-pt scale "
        f"rescaled to −2 (very negative) … +2 (very positive). The table "
        f"below shows every factor and the change from {first_y} to {last_y}, "
        f"sorted by absolute change. The figure shows the top 10 movers.",
        italic=True, size=10, color=GRAY_RGB,
    )
    if p_factor_movers:
        add_image(p_factor_movers, width_in=6.5)
        add_para("Figure 7. Top 10 job-factor movers — grouped by year.",
                 italic=True, size=10, color=GRAY_RGB)
    # Full factor table sorted by absolute delta
    rows = []
    for label, fm, lm, d in sorted(factor_delta_pairs,
                                   key=lambda r: abs(r[3]),
                                   reverse=True):
        row = [label]
        for y in years:
            fi = {f["factor"]: f for f in results_by_year[y]["factors"]}
            m = (fi.get(label) or {}).get("mean")
            row.append(f"{m:+.2f}" if m is not None else "—")
        row.append(_fmt_delta(d))
        rows.append(row)
    headers = ["Job factor"] + [str(y) for y in years] + [delta_col]
    widths = [Inches(3.2)] + [Inches(0.7)] * len(years) + [Inches(0.7)]
    make_table(headers, rows, col_widths=widths)
    doc.add_paragraph()

    # ----- 9. Leadership -----
    add_heading("8. Leadership Item Changes", 1)
    add_para(
        f"Leadership items scored 1–5 (5 = strongly agree). Table shows "
        f"every leadership item with year-by-year means and the {first_y}→"
        f"{last_y} delta, sorted by absolute change. Figure shows the top "
        f"8 movers.",
        italic=True, size=10, color=GRAY_RGB,
    )
    if p_ld_movers:
        add_image(p_ld_movers, width_in=6.5)
        add_para("Figure 8. Top 8 leadership-item movers — grouped by year.",
                 italic=True, size=10, color=GRAY_RGB)
    rows = []
    for label, fm, lm, d in sorted(ld_delta_pairs,
                                   key=lambda r: abs(r[3]),
                                   reverse=True):
        row = [label]
        for y in years:
            ld_idx = {it["item"]: it for it in ld_items_by_year[y]}
            it = ld_idx.get(label) or {}
            m = it.get("mean")
            row.append(f"{m:.2f}" if m is not None else "—")
        row.append(_fmt_delta(d))
        rows.append(row)
    headers = ["Leadership item"] + [str(y) for y in years] + [delta_col]
    widths = [Inches(3.4)] + [Inches(0.7)] * len(years) + [Inches(0.7)]
    make_table(headers, rows, col_widths=widths)
    doc.add_paragraph()

    # ----- Footer -----
    add_para(
        "Generated from REDCap exports across loaded survey years. "
        "REDCap responses are anonymous; longitudinal comparisons are "
        "aggregate only.",
        italic=True, size=9, color=GRAY_RGB,
    )

    # Save
    if isinstance(output_path, (str, os.PathLike)):
        doc.save(output_path)
    else:
        doc.save(output_path)
